from mysql2doc.MySql import MySql

__author__ = 'yanglikun'
from docx import Document


class Word:
    def __init__(self, document):
        super().__init__()
        self.document = document

    def __bold(ele, text):
        ele.paragraphs[0].add_run(text).bold = True

    def addTable(self, table, seqNO):
        self.document.add_heading("{} {}".format(seqNO, table.name), level=1)
        self.document.add_paragraph(table.comment)
        tableGrid = self.document.add_table(rows=1, cols=4, style='TableGrid')
        titleRow = tableGrid.rows[0].cells
        Word.__bold(titleRow[0], '字段')
        Word.__bold(titleRow[1], '类型')
        Word.__bold(titleRow[2], '备注')
        Word.__bold(titleRow[3], '允许为空')
        for field in table.fields:
            row = tableGrid.add_row().cells
            row[0].text = field.name
            row[1].text = field.type
            row[2].text = field.comment
            if field.nullable:
                row[3].text = '是'
            else:
                row[3].text = '否'
        self.document.add_paragraph()
        self.document.add_paragraph('索引列', style='ListBullet')
        idxTableGrid = self.document.add_table(rows=1, cols=5, style='TableGrid')
        idxTitleRow = idxTableGrid.rows[0].cells
        Word.__bold(idxTitleRow[0], '唯一索引')
        Word.__bold(idxTitleRow[1], '索引名称')
        Word.__bold(idxTitleRow[2], '索引顺序')
        Word.__bold(idxTitleRow[3], '字段')
        Word.__bold(idxTitleRow[4], '备注')
        for index in table.indices:
            idxRow = idxTableGrid.add_row().cells
            if index.isUnique:
                Word.__bold(idxRow[0], '是');
            else:
                Word.__bold(idxRow[0], '否');
            idxRow[1].text = index.name
            idxRow[2].text = index.seqNO
            idxRow[3].text = index.fieldName
            idxRow[4].text = index.comment
        pass

    def createFile(fileName='database'):
        document = Document()
        document.add_heading('数据库表结构', 0)
        document.add_paragraph('数据库表结构')
        mysql=None
        try:
            mysql = MySql()
            word = Word(document)
            for idx, table in enumerate(mysql.generateTableData(), start=1):
                word.addTable(table, idx)
        finally:
            MySql.close(mysql)
        document.add_page_break()
        document.save(fileName + '.docx')
